"""Wordテンプレートへのデータ書き込みモジュール"""
from docx import Document
from docx.shared import Pt
from pathlib import Path
from datetime import datetime
from typing import Optional
import re

from .company_form import CompanyInfo, SocialInsuranceLabor
from .upload_handler import CurriculumGroup


class WordWriter:
    """Wordテンプレートへの書き込みを行うクラス"""

    def __init__(self, template_dir: str, output_dir: str):
        self.template_dir = Path(template_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _load_template(self, template_path: str) -> Document:
        """テンプレートファイルを読み込む"""
        full_path = self.template_dir / template_path
        return Document(str(full_path))

    def _save_document(self, doc: Document, filename: str) -> str:
        """ドキュメントを保存する"""
        # ファイル名をサニタイズ
        filename = self._sanitize_filename(filename)
        output_path = self.output_dir / filename
        doc.save(str(output_path))
        return str(output_path)

    def _sanitize_filename(self, filename: str) -> str:
        """ファイル名から不正な文字を除去"""
        # タブ、改行、制御文字を除去
        filename = re.sub(r'[\t\n\r\x00-\x1f\x7f]', '', filename)
        # Windowsで使用できない文字を置換
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        # 前後の空白を除去
        filename = filename.strip()
        return filename if filename else "unnamed.docx"

    def _replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str):
        """段落内のテキストを置換する"""
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

    def _replace_text_in_document(self, doc: Document, replacements: dict):
        """ドキュメント全体でテキストを置換する"""
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    self._replace_text_in_paragraph(paragraph, old_text, new_text)

        # テーブル内も置換
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                self._replace_text_in_paragraph(paragraph, old_text, new_text)

    def write_事業内職業能力開発計画(self, company: CompanyInfo,
                                    create_date: datetime) -> str:
        """11_事業内職業能力開発計画.docxを作成"""
        doc = self._load_template("計画申請/11_事業内職業能力開発計画.docx")

        # 置換するテキスト
        replacements = {
            "株式会社○○○○": company.company_name,
            "〇〇〇〇": company.main_business if company.main_business else "（経営方針・理念を記載）",
            "年　　月　　日作成": f"{create_date.year}年{create_date.month}月{create_date.day}日作成",
        }

        self._replace_text_in_document(doc, replacements)

        filename = f"11_事業内職業能力開発計画_{company.company_name}.docx"
        return self._save_document(doc, filename)

    @staticmethod
    def _fill_reiwa_date_paragraph(para, year: int, month: int) -> bool:
        """段落内の「令和 ? 年 ? 月 ? 日」形式の日付を置換する。
        年・月の数字を埋め、日 (digit + ' 日' ラベル含む) を空欄にする。

        Returns True if a date was found and filled, False otherwise.
        """
        runs = para.runs
        # 確認: 「令和」「年」「月」を含むか
        full_text = ''.join(r.text for r in runs)
        if '令和' not in full_text or '年' not in full_text or '月' not in full_text:
            return False

        # 「令和」「年」「月」「日」ラベル直後の数字を順次置換
        # ラベルは run.text に '令和' / '年' / '月' / '日' を含む形で出現する。
        expecting = None  # 'year' | 'month' | 'day'
        day_clear_indices: list[int] = []

        for i, run in enumerate(runs):
            t = run.text
            stripped = t.replace('　', '').replace(' ', '').strip()
            if '日' in t:
                # 「日」ラベル run — 日数字 + ラベル両方をクリア
                day_clear_indices.append(i)
                expecting = None
                continue
            if '令和' in t:
                expecting = 'year'
                continue
            if '年' in t:
                expecting = 'month'
                continue
            if '月' in t:
                expecting = 'day'
                continue
            # 数字 run の処理
            if stripped.isdigit():
                if expecting == 'year':
                    run.text = str(year - 2018)  # 令和N = year - 2018
                    expecting = None
                elif expecting == 'month':
                    run.text = str(month)
                    expecting = None
                elif expecting == 'day':
                    day_clear_indices.append(i)
            elif expecting == 'day' and stripped == '':
                # 「月」と日数字の間の空白
                day_clear_indices.append(i)

        # 日関連 run をすべてクリア
        for i in day_clear_indices:
            runs[i].text = ''

        return True

    def write_提出代行証明書_アスラク(self, sr: SocialInsuranceLabor,
                                     submit_date: datetime) -> str:
        """提出代行に関する証明書 ※アスラク.docx を作成
        上部の「令和 N 年 M 月 D 日」を「令和 N 年 M 月」に自動入力 (日は空欄)
        """
        template_path = "計画申請/提出代行に関する証明書 ※アスラク.docx"
        doc = self._load_template(template_path)

        for para in doc.paragraphs:
            if self._fill_reiwa_date_paragraph(para, submit_date.year, submit_date.month):
                break

        filename = "提出代行に関する証明書 ※アスラク.docx"
        return self._save_document(doc, filename)

    def generate_word_documents(self, company: CompanyInfo,
                                 group: CurriculumGroup,
                                 create_date: datetime,
                                 sr: Optional[SocialInsuranceLabor] = None) -> list:
        """Word書類一式を生成。テンプレートが無いものはスキップする (書式別の互換用)"""
        generated_files = []

        # 11_事業内職業能力開発計画 (現行書式のみ)
        if (self.template_dir / "計画申請/11_事業内職業能力開発計画.docx").exists():
            try:
                file_path = self.write_事業内職業能力開発計画(company, create_date)
                generated_files.append(file_path)
            except Exception as e:
                print(f"事業内職業能力開発計画の生成に失敗: {e}")

        # 提出代行に関する証明書 ※アスラク
        if (self.template_dir / "計画申請/提出代行に関する証明書 ※アスラク.docx").exists():
            try:
                file_path = self.write_提出代行証明書_アスラク(sr, create_date)
                generated_files.append(file_path)
            except Exception as e:
                print(f"提出代行証明書(アスラク)の生成に失敗: {e}")

        return generated_files
